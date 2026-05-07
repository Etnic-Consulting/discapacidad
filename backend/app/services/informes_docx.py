"""
L16 · Constructor de informes en formato Word desde JSON canónico + JSON LLM.

NO convierte HTML · arma estructura nativa Word con estilos institucionales.
"""
from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PRIMARY = RGBColor(0x01, 0x4a, 0x30)
GREEN_MID = RGBColor(0x2e, 0x7d, 0x4f)
INK_MUTED = RGBColor(0x55, 0x55, 0x55)

TITULOS = {
    "pueblo": "PUEBLO",
    "dpto": "DEPARTAMENTO",
    "mpio": "MUNICIPIO",
    "resguardo": "RESGUARDO",
    "macro": "MACRORREGIÓN",
}


def fmt(n) -> str:
    if n is None:
        return "—"
    if isinstance(n, float):
        return f"{n:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")
    if isinstance(n, int):
        return f"{n:,}".replace(",", ".")
    return str(n)


def _safe_get(d, *keys, default=None):
    cur = d
    for k in keys:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(k)
    return cur if cur is not None else default


def _add_h1(doc: Document, txt: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = PRIMARY


def _add_h2(doc: Document, txt: str):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = PRIMARY


def _add_h3(doc: Document, txt: str):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = GREEN_MID


def _add_meta(doc: Document, txt: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = INK_MUTED


def build_docx(data: dict[str, Any], llm: dict[str, Any] | None, output_path: str) -> None:
    """Construye el .docx desde JSON canónico + JSON LLM."""
    if llm is None:
        llm = {}

    nivel = data.get("tipo", "pueblo")
    nombre = data.get("nombre", "—")
    fecha = data.get("fecha_generacion", "")[:10]
    periodo = data.get("periodo", "2018")
    sec = data.get("secciones", {})

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    # ============== PORTADA ==============
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ORGANIZACIÓN NACIONAL INDÍGENA DE COLOMBIA — ONIC")
    r.bold = True
    r.font.color.rgb = PRIMARY
    r.font.size = Pt(11)

    _add_meta(doc, "Sistema de Monitoreo Territorial · Módulo Capacidades Diversas")
    doc.add_paragraph()
    doc.add_paragraph()

    _add_h1(doc, f"{TITULOS.get(nivel, 'INFORME').upper()} {nombre}")
    _add_meta(doc, f"Informe territorial · Personas con Capacidades Diversas")
    _add_meta(doc, f"CNPV {periodo} · Generado: {fecha}")
    doc.add_page_break()

    # ============== CRÉDITOS INSTITUCIONALES ==============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ORGANIZACIÓN NACIONAL INDÍGENA DE COLOMBIA")
    r.bold = True
    r.font.color.rgb = PRIMARY
    r.font.size = Pt(14)

    _add_meta(doc, "Sistema de Monitoreo Territorial · SMT-ONIC")
    doc.add_paragraph()

    _add_h3(doc, "Consejería Mayor de Gobierno")
    consejeros = [
        ("Consejero Mayor", "ROSELINO GUARUPE JOROPA"),
        ("Consejería Secretaría General", "GUILLIANA PATRICIA ARRIETA MAURY"),
        ("Consejería de Planeación, Administración y Finanzas", "Sandra Dulfay Perdomo Prieto"),
        ("Consejería de Sistemas de Investigación, Información y Comunicaciones", "Milton Piranga Cruz"),
        ("Consejería de Derechos de los Pueblos Indígenas, Derechos Humanos y Paz", "RIDER PAY NASTACÚAS"),
        ("Consejería de Educación Propia e Intercultural", "Yeferson David Domicó"),
        ("Consejería de Mujer, Familia y Generación", "Dora Liseth Garcés Aguablanca"),
        ("Consejería de Planes de Vida y Desarrollo Propio", "Verónica Solís Fuentes"),
        ("Consejería de Medicina Tradicional y Salud Occidental", "Johny Jefferson Ramírez"),
        ("Consejería de Territorio, Recursos Naturales y Biodiversidad", "arlenys ester alvarado epiayu"),
    ]
    for cargo, nom in consejeros:
        p = doc.add_paragraph()
        rc = p.add_run(cargo + "\n")
        rc.font.size = Pt(9)
        rc.font.color.rgb = GREEN_MID
        rn = p.add_run(nom)
        rn.bold = True
        rn.font.size = Pt(11)

    doc.add_paragraph()
    _add_h3(doc, "Coordinación Técnica")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Wilson Fernando Herrera Baltán\n")
    r.bold = True
    p.add_run("Coordinador Técnico Sistema de Monitoreo Territorial SMT-ONIC\n").italic = True
    p.add_run("\n")
    r2 = p.add_run("Eudo Fidel Cuarán Hernández\n")
    r2.bold = True
    p.add_run("Coordinador Social").italic = True

    doc.add_page_break()

    # ============== §2 RESUMEN EJECUTIVO ==============
    _add_h2(doc, "§ 2. Resumen ejecutivo")
    for parr in str(llm.get("ejecutivo", "Análisis pendiente.")).split("\n"):
        if parr.strip():
            doc.add_paragraph(parr.strip())

    # KPIs cabecera
    pob = _safe_get(sec, "demografia", "poblacion_total", "value")
    con_disc = _safe_get(sec, "capacidades_diversas", "con_discapacidad", "value")
    prev = _safe_get(sec, "capacidades_diversas", "prevalencia_x_1000", "value")
    n_dptos = _safe_get(sec, "territorial", "n_departamentos")
    n_resg = _safe_get(sec, "territorial", "n_resguardos")
    icv_score = _safe_get(sec, "icv", "score")

    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Población"
    hdr[1].text = "Cap. diversas"
    hdr[2].text = "Territorios"
    hdr[3].text = "ICV"
    row = t.add_row().cells
    row[0].text = fmt(pob)
    row[1].text = f"{fmt(con_disc)} ({fmt(prev)}‰)"
    row[2].text = f"{n_dptos or '—'} dptos · {n_resg or '—'} resg"
    row[3].text = str(icv_score) if icv_score is not None else "—"

    # ============== §3 DEMOGRAFÍA ==============
    _add_h2(doc, "§ 3. Demografía")
    doc.add_paragraph(str(llm.get("demografia", "Análisis pendiente.")))

    # ============== §4 CAPACIDADES DIVERSAS ==============
    _add_h2(doc, "§ 4. Capacidades diversas")
    doc.add_paragraph(str(llm.get("prevalencia", "Análisis pendiente.")))

    cd = sec.get("capacidades_diversas", {})
    tipos = cd.get("tipos_limitacion", [])[:5]
    if tipos:
        _add_h3(doc, "Tipos de limitación (top 5)")
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light List Accent 1"
        t.rows[0].cells[0].text = "Tipo"
        t.rows[0].cells[1].text = "Total"
        for tp in tipos:
            row = t.add_row().cells
            row[0].text = str(tp.get("tipo_limitacion", ""))
            row[1].text = fmt(tp.get("total", 0))

    # ============== §5 TERRITORIAL ==============
    _add_h2(doc, "§ 5. Distribución territorial")
    doc.add_paragraph(str(llm.get("territorial", "Análisis pendiente.")))

    dptos = _safe_get(sec, "territorial", "departamentos", default=[]) or []
    if dptos:
        _add_h3(doc, "Departamentos con mayor presencia")
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light List Accent 1"
        for i, h in enumerate(["Departamento", "Población", "Con cap. diversas", "Prevalencia"]):
            t.rows[0].cells[i].text = h
        for d in dptos[:8]:
            row = t.add_row().cells
            row[0].text = str(d.get("nom_dpto", "—"))
            row[1].text = fmt(d.get("total"))
            row[2].text = fmt(d.get("con_discapacidad"))
            tasa = d.get("tasa_x_1000")
            row[3].text = f"{fmt(float(tasa) if tasa else None)}‰"

    # ============== §6 LENGUA + §7 NBI ==============
    _add_h2(doc, "§ 6. Lengua y cultura")
    doc.add_paragraph(str(llm.get("lengua", "Análisis pendiente.")))

    _add_h2(doc, "§ 7. Condiciones de vida (NBI)")
    doc.add_paragraph(str(llm.get("vida", "Análisis pendiente.")))
    nbi = sec.get("nbi", {})
    if "pct_nbi" in nbi:
        pct = _safe_get(nbi, "pct_nbi", "value")
        prom = nbi.get("promedio_indigena_nbi", 67.6)
        doc.add_paragraph(f"NBI {nombre}: {fmt(pct)}% · Promedio indígena nacional: {fmt(prom)}%")

    # ============== §8 CONFLICTO + §9 ICV ==============
    _add_h2(doc, "§ 8. Conflicto armado")
    doc.add_paragraph(str(llm.get("conflicto", "Análisis pendiente.")))
    conf = sec.get("conflicto", {})
    n_vict = _safe_get(conf, "victimas_total_disc", "value")
    hp = conf.get("hecho_principal") or {}
    if n_vict:
        doc.add_paragraph(f"Víctimas con cap. diversas: {fmt(n_vict)} · Hecho principal: {hp.get('hecho', '—')} ({fmt(conf.get('pct_hecho_principal'))}%)")

    _add_h2(doc, "§ 9. Indicador Compuesto de Vulnerabilidad (ICV)")
    doc.add_paragraph(str(llm.get("icv", "Análisis pendiente.")))
    icv = sec.get("icv", {})
    if icv and "componentes" in icv:
        c = icv["componentes"]
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light List Accent 1"
        for i, h in enumerate(["Componente", "Normalizado", "Aporte"]):
            t.rows[0].cells[i].text = h
        for nom, key, peso in [("NBI", "nbi_norm", 0.4), ("Prevalencia", "prevalencia_norm", 0.3), ("Conflicto", "conflicto_norm", 0.3)]:
            row = t.add_row().cells
            row[0].text = nom
            v = c.get(key, 0)
            row[1].text = str(v)
            row[2].text = f"{v * peso:.3f}"
        row = t.add_row().cells
        row[0].text = "ICV TOTAL"
        row[2].text = str(icv.get("score", "—"))

    # ============== §10 RECOMENDACIONES ==============
    _add_h2(doc, "§ 10. Recomendaciones de política pública")
    recs = llm.get("recomendaciones", []) or []
    if isinstance(recs, str):
        recs = [r.strip() for r in recs.split("\n") if r.strip()]
    if not recs:
        recs = ["Recomendaciones pendientes."]
    for i, r in enumerate(recs, 1):
        doc.add_paragraph(f"{i}. {r}")

    # ============== §11 METODOLOGÍA ==============
    _add_h2(doc, "§ 11. Metodología")
    doc.add_paragraph(
        "Fuentes: Censo Nacional de Población y Vivienda 2018 (DANE) · "
        "Registro Único de Víctimas (UARIV) · MGN 2025 DANE · catálogos canónicos ONIC."
    )
    doc.add_paragraph(f"Período: {periodo} para indicadores demográficos. 1985-2024 para registro RUV.")
    doc.add_paragraph(
        "Trazabilidad: cada cifra está taggeada con metadata "
        "{value, query, table, period, n, confiabilidad} consultable vía endpoint "
        f"GET /api/v1/informes/{nivel}/{data.get('id', '—')}/data"
    )
    doc.add_paragraph()
    _add_meta(doc, f"SMT-ONIC · Sistema de Monitoreo Territorial · {fecha}")
    _add_meta(doc, "poblacion@onic.org.co · Datos abiertos CC BY 4.0")

    doc.save(output_path)
